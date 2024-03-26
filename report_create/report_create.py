import argparse
import json
import docx
import os.path
import ctypes
import traceback
import datetime
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from fuzzywuzzy import fuzz

myappid = 'mycompany.myproduct.subproduct.version'  # arbitrary string
ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)


def section1(file, table):
    """ Section 1: General Information """
    section_titles = ('Name', 'Domain', 'UserName')  # Titles in the template to be filled
    row = table.rows[2]  # Row in the table
    for i in range(len(section_titles)):  # Iterating through the sections
        if section_titles[i] in file:  # If the section exists in the JSON file
            row.cells[i].text = str(file[section_titles[i]])  # Add the value from the dictionary to the template
        else:
            continue


def section2(file, table, os_version):
    """ Section 2: Operating System """
    section_titles = ('Caption', 'OSArchitecture', 'Version')  # Sections in the template to be filled
    row = table.rows[2]  # Row in the table
    for i in range(len(section_titles)):  # Iterating through the sections
        if section_titles[i] in file:  # If the section exists in the JSON file
            row.cells[i].text = str(file[section_titles[i]])  # Add the value from the dictionary to the template
            if i == 2 and 'BuildNumber' in file:
                if os_version != str(file['BuildNumber']):  # If the OS version is outdated
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))  # Highlight the cell
                    row.cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
        else:
            continue


def section3(file, table):
    """ Section 3: Network Configuration """
    section_titles = ('ip', 'gateway', 'dns', 'mac')  # Sections in the template to be filled
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, adapter in enumerate(sorted(file)):  # Iterating through adapters
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Getting cells for the current row
        row_cells[0].text = str(adapter)  # Adding adapter name to the template
        for j in range(len(section_titles)):  # Iterating through sections
            for value in file[adapter]:  # Iterating through values
                if section_titles[j] in value:  # If the section exists in the JSON file
                    row_cells[j + 1].text = str(
                        value[section_titles[j]])  # Adding the value from the dictionary to the template
                    break
                else:
                    continue
            else:
                continue


def section4(file, table):
    """ Section 4: Time Service """
    section_titles = ('State', 'NtpSource', 'TimeZone', 'Last_time')  # Sections in the template to be filled
    row = table.rows[2]  # The row in the table
    for i in range(len(section_titles)):  # Iterating through the sections
        if section_titles[i] in file:  # If the section exists in the JSON file
            row.cells[i].text = str(file[section_titles[i]])  # Add the value from the dictionary to the template
        else:
            continue


def section5(file, table):
    """ Section 5: OS Updates """
    row = table.rows[2]  # The row in the table
    hotfixes_text = ''  # A string to accumulate all updates
    for key in file:  # Iterating through the entries
        if key != 'Other':  # If it's an OS update
            hotfixes_text += f'{key} from {str(file[key][0])}\n'  # Add the update number and date to the string
        else:
            row.cells[0].text = str(file[key][0])  # Add the update service property to the template
            row.cells[1].text = str(file[key][1])  # Add another update service property to the template
            row.cells[3].text = str(file[key][2])  # Add another update service property to the template
    row.cells[2].text = hotfixes_text.strip()  # Add all updates to the cell


def section6(file, table):
    """ Section 6: User Accounts """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, account in enumerate(sorted(file)):  # Iterating through user accounts
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = str(account)  # Add the account name to the template
        value = file[account]  # Get the account properties
        for j in range(len(value)):  # Iterate through account properties
            if j == 0:  # If it's the account status
                if value[j]:
                    row_cells[1].text = 'Отключен'
                    for cell in row_cells:  # If the account is inactive, shade all cells in the row
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DDDDDD"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_elm_1)
                else:
                    row_cells[1].text = 'Активен'
            elif j == 1:  # If it's Admin status
                row_cells[j + 1].text = str(value[j])  # Write the value
                if value[j] in ('Администраторы', 'Administrators') and not value[0]:
                    for cell in row_cells:  # If the account is an admin and active, shade all cells in the row
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                        cell._tc.get_or_add_tcPr().append(shading_elm_1)
            else:
                row_cells[j + 1].text = str(value[j])  # Add the rest of the section values


def section7(file, table, std_groups):
    """ Section 7: Groups """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, group in enumerate(sorted(file)):  # Iterating through groups
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = str(group)  # Add the group names to the template
        if file[group]:  # If there are members in the group, add them to the template
            row_cells[1].text = '\n'.join(file[group])  # Each group member on a new line
        if group in std_groups:
            row_cells[2].text = 'Да'  # The group is standard
        else:
            row_cells[2].text = 'Нет'  # The group is non-standard
            for cell in row_cells:  # Shade all cells in the row if the group is non-standard
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_1)


def section8(file, table):
    """ Section 8: Shared Folders """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, share in enumerate(sorted(file)):  # Iterating through shared folders in the JSON file
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = str(share)  # Add the shared folder name to the template
        row_cells[1].text = str(file[share][0])  # Add the folder path to the template
        row_cells[2].text = str(file[share][1])  # Add the description to the template


def section9(file, table):
    """ Section 9: Antivirus Protection """
    def write_avz_table(avz_name, avz_version, avz_server, avz_base, tb_cells, j):
        row_cells = tb_cells[j * columns: (j + 1) * columns]  # Get cells for the current row
        row_cells[0].text = str(avz_name)  # Add the avz name
        row_cells[1].text = str(avz_version)  # Add the avz version
        row_cells[2].text = str(avz_server)  # Add the avz server
        row_cells[3].text = str(avz_base)  # Add the date of the antivirus definitions
        if avz_base:  # If the date exists
            # Convert the definitions date (year, month, day)
            date_avz = list(map(int, avz_base.split('.')[::-1]))
            # Calculate the difference between dates
            if (data - datetime.date(date_avz[0], date_avz[1], date_avz[2])).days > 1:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))  # Highlight the cell
                row_cells[3]._tc.get_or_add_tcPr().append(shading_elm_1)

    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    data = datetime.date.today()  # Today's date
    if isinstance(file, dict):
        write_avz_table(file['Name'], file['Version'], file['Server'], file['Base'], table_cells, 0)
    else:
        for i, avz in enumerate(file):  # Iterating through avz in the JSON file
            write_avz_table(avz['Name'], avz['Version'], avz['Server'], avz['Base'], table_cells, i)


def section10(file, table, folder, match_ratio):
    """ Section 10: Software """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    with open(f'{folder}\\src\\data\\software.json', 'rb') as soft_file:  # Open the list of software
        std_soft = json.load(soft_file)  # Import the JSON file
    for i, application in enumerate(file):  # Iterating through applications in the JSON file
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = f'{i + 1}.'  # Add the serial number
        row_cells[1].text = str(application['DisplayName'])  # Add the standard software name
        row_cells[2].text = str(application['DisplayVersion'])  # Add the standard software version
        match_found = False  # Flag for match found
        for soft in std_soft:  # Iterate through all software in the list
            ratio = fuzz.partial_ratio(application['DisplayName'], soft)  # Calculate the match percentage for names
            if ratio > match_ratio:
                match_found = True
                row_cells[3].text = f'Да,\nВерсия: {str(std_soft[soft])}'
                break
            else:
                continue
        if not match_found:  # If no matches found
            row_cells[3].text = 'Нет'
            for cell in row_cells:  # Highlight all cells in the row if no match is found
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_1)


def section11(file, table):
    """ Section 11: Password Policy """
    section_titels = [str(j) for j in range(8)]  # Titles for sections to be filled in the template
    for i in range(len(section_titels)):  # Iterating through the sections
        if section_titels[i] in file:  # If the section exists in the JSON file
            row = table.rows[i + 1]  # The row in the table
            row.cells[1].text = str(file[section_titels[i]])  # Add the value from the dictionary to the template
        else:
            continue


def section13(file, table):
    """ Section 13: Autostart Applications """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, autorun_app in enumerate(sorted(file)):  # Iterating through autostart applications in the JSON file
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = f'{i + 1}.'  # Add the serial number
        row_cells[1].text = str(autorun_app)  # Add the name of the autostart application
        row_cells[2].text = str(file[autorun_app])  # Add the path of the application


def section14(file, table):
    """ Section 14: Task Scheduler """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, task in enumerate(sorted(file)):  # Iterating through tasks in the JSON file
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = f'{i + 1}.'  # Add the serial number
        row_cells[1].text = str(task)  # Add the name of the task from the scheduler
        row_cells[2].text = str(file[task]['State'])  # Add the status of the task
        row_cells[3].text = str(file[task]['Author'])  # Add the author of the task


def section15(file, table, product_type):
    """ Section 15: System Software """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    if product_type == 1:
        for i, sys_application in enumerate(file):  # Iterating through system applications in the JSON file
            row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
            row_cells[0].text = f'{i + 1}.'  # Add the serial number
            row_cells[1].text = str(sys_application['Name'])  # Add the name of the standard software
            row_cells[2].text = str(sys_application['Version'])  # Add the version of the standard software
    else:
        for i, role in enumerate(file):  # Iterating through roles in the JSON file
            row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
            row_cells[0].text = f'{i + 1}.'  # Add the serial number
            row_cells[1].text = str(role['Name'])  # Add the role name
            row_cells[2].text = str(role['InstallState'])  # Add the role's installation state
            row_cells[3].text = str(role['Description'])  # Add the role description


def section16(file, table, folder):
    """ Section 16: Services """
    rename = {
        'Disabled': 'Отключена',
        'Manual': 'Вручную',
        'Auto': 'Автоматически',
        'Stopped': 'Остановлена',
        'Running': 'Выполняется',
        'Unknown': 'Неизвестно'
    }
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    with open(f'{folder}\\src\\data\\services.json', 'rb') as srvc_file:  # Open the list of services
        services = json.load(srvc_file)  # Import the JSON file
    diff_services = set([x['Name'] for x in file]) - set(
        [x['Name'] for x in services])  # Differences from the list of standard services
    for i, service in enumerate(file):  # Iterating through services
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = f'{i + 1}.'  # Add the serial number
        row_cells[1].text = str(service['DisplayName'])  # Add the service name
        if service['Startmode'] in rename:  # If the service state is in the dictionary, translate it
            row_cells[2].text = str(rename[service['Startmode']])  # Add the startup mode of the service
        else:
            row_cells[2].text = 'Неизвестно'
        if service['State'] in rename:
            row_cells[3].text = str(rename[service['State']])  # Add the service status
        else:
            row_cells[3].text = 'Неизвестно'
        # Highlight all cells in the row if the service is non-standard or running
        if service['Name'] in diff_services:
            for cell in row_cells:
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm_1)
        else:
            if row_cells[3].text == 'Выполняется':
                for cell in row_cells:
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DDDDDD"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm_1)

        row_cells[4].text = str(service['Name'])  # Add the service's system name
        row_cells[5].text = str(service['PathName'])  # Add the service's path


def section17(file, table):
    """ Section 17: System Log """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, system_log in enumerate(sorted(file, key=lambda x: int(x))):  # Iterating through entries in the JSON file
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = str(system_log)  # Add the event ID from the log
        row_cells[1].text = str(file[system_log]['Time'])  # Add the event date
        row_cells[2].text = str(file[system_log]['ProviderName'])  # Add the event source
        row_cells[3].text = str(file[system_log]['Level'])  # Add the event level
        row_cells[4].text = str(file[system_log]['RecordId'])  # Add the event record ID
        row_cells[5].text = str(file[system_log]['Message'])  # Add the event description


def section18(file, table):
    """ Section 18: Security Log """
    columns = len(table.columns)  # Number of columns
    table_cells = table._cells[columns * 2:]  # All table cells, skipping the first two rows
    for i, security_log in enumerate(
            sorted(file, key=lambda x: int(x), reverse=True)):  # Iterating through entries in the JSON file
        row_cells = table_cells[i * columns: (i + 1) * columns]  # Get cells for the current row
        row_cells[0].text = str(security_log)  # Add the serial number of the event from the log
        row_cells[1].text = str(file[security_log]['Time'])  # Add the date of the event
        row_cells[2].text = str(file[security_log]['Id'])  # Add the event ID
        # Add the event description, trimming at the first carriage return
        row_cells[3].text = str(
            file[security_log]['Message'][:file[security_log]['Message'].find('\r')])


def parse_arguments():
    """ Parses command-line arguments provided to the script. """
    parser = argparse.ArgumentParser(description="This script analyzes computer parameters from JSON files and "
                                                 "generates a report in DOCX format.")

    parser.add_argument("-c", "--computer", required=True,
                        help="Name of the computer for which the report is generated.")
    parser.add_argument("-p", "--path", required=True,
                        help="Root directory where the CheckARM.exe executable file is located.")
    parser.add_argument("-rv", "--report_version", required=True, choices=["short", "full"],
                        help="Report version. Can be 'short' or 'full'. The script selects a specific report template "
                             "and the number of blocks generated based on this version.")
    parser.add_argument("-ov", "--os_version", required=True,
                        help="Default operating system version. In the report block 2 'Operating System', the script "
                             "compares the specified version with the version of the OS of the computer being checked.")
    parser.add_argument("-r", "--ratio", type=int, default=90,
                        help="Ratio used for fuzzy string matching of software names by the FuzzyWuzzy library. "
                             "Defaults to 90.")
    args = parser.parse_args()
    return args


def report_create(computer, folder, report_version, os_version, ratio):
    """
    The function `report_create` generates a report based on specified sections and saves it as a Word
    document.
    
    :param computer: The `computer` parameter in the `report_create` function represents the name or
    identifier of the computer for which the report is being generated. It is used to create paths and
    filenames specific to that computer within the report generation process
    :param folder: The `folder` parameter in the `report_create` function represents the directory path
    where the report files and related service files are stored. It is used to specify the location of
    JSON files, templates, reports, and error logs within the file system
    :param report_version: The `report_version` parameter in the `report_create` function is used to
    specify whether a short report or a full report should be generated. If the value of
    `report_version` is set to `'short'`, the function will generate a short report. Otherwise, if the
    value is `'full'`
    :param os_version: The `os_version` parameter in the `report_create` function is used to pass the
    operating system version of the computer for which the report is being generated. It is used within
    the function to populate specific sections of the report based on the operating system version
    provided
    :param ratio: The `ratio` parameter in the `report_create` function is used for fuzzy string matching
    of software names by the FuzzyWuzzy library. Default value is 90."
    :return: The function `report_create` is returning the file path of the generated report in the
    format `{folder}\reports\{computer}_{dt_now}.docx` after creating the report. If an error occurs
    during the process, it will log the error in the specified log file and return the string 'Error'.
    """
    i = 'report_create'
    try:
        json_path = f'{folder}\\src\\jsons\\{computer}\\'
        # Dictionary specifying section numbers, file paths, and function names
        sections_dict = {
            1: (f'{json_path}1_comp.json', section1),
            2: (f'{json_path}2_os.json', section2),
            3: (f'{json_path}3_net.json', section3),
            4: (f'{json_path}4_time.json', section4),
            5: (f'{json_path}5_hotfixes.json', section5),
            6: (f'{json_path}6_accounts.json', section6),
            7: (f'{json_path}7_groups.json', section7),
            8: (f'{json_path}8_shares.json', section8),
            9: (f'{json_path}9_avz.json', section9),
            10: (f'{json_path}10_1_std_software.json', section10),
            11: (f'{json_path}16_pass.json', section11),
            13: (f'{json_path}11_autorun.json', section13),
            14: (f'{json_path}12_tasks.json', section14),
            15: (f'{json_path}10_2_sys_software.json', section15),
            16: (f'{json_path}14_services.json', section16),
            17: (f'{json_path}13_1_systemlog.json', section17),
            18: (f'{json_path}13_2_securitylog.json', section18)
        }
        if os.path.exists(f"{folder}\\src\\jsons\\{computer}\\2_os.json"):
            with open(f"{folder}\\src\\jsons\\{computer}\\2_os.json", 'rb') as json_file:
                product_type = json.load(json_file)['ProductType']
        else:
            product_type = 1
        if product_type == 1:
            template = f"{folder}\\src\\reporting\\templates\\КЛ АРМ_шаблон"
        else:
            template = f"{folder}\\src\\reporting\\templates\\КЛ СРВ_шаблон"
        if report_version == 'short':  # If a short report is needed
            doc_template = docx.Document(f"{template}_short.docx")  # Open the template for the short report
            sections = (1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11)
        else:
            doc_template = docx.Document(f"{template}_full.docx")  # Open the template for the full report
            sections = (1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 13, 14, 15, 16, 17, 18)
        for i in sections:  # Range of sections
            table = doc_template.tables[i]  # Select the appropriate table
            path = sections_dict[i][0]
            if os.path.exists(path):  # Check if the file exists and is not empty
                if os.path.getsize(path) != 0:
                    with open(path, 'rb') as json_file:
                        file = json.load(json_file)  # Import the JSON file
                        if i in (1, 2, 4, 5, 11):  # Sections that can have only one row
                            if i == 2:
                                table.add_row()  # Add a row
                                section2(file, table, os_version)
                            elif i == 11:
                                sections_dict[i][1](file, table)
                            else:
                                table.add_row()  # Add a row
                                sections_dict[i][1](file, table)
                        else:  # Sections that can have more than one row
                            if i == 9 and isinstance(file, dict):
                                table_len = 1
                            else:
                                table_len = len(file)
                            for _ in range(table_len):  # If more than one row is needed
                                table.add_row()  # Add a row
                            if i == 7:  # If it's the groups section, read the list of standard groups
                                with open(f'{folder}\\src\\data\\groups.json', 'rb') as groups_json:
                                    std_groups = json.load(groups_json)['groups']
                                sections_dict[i][1](file, table, std_groups)
                            elif i == 10:
                                sections_dict[i][1](file, table, folder, ratio)
                            elif i == 16:
                                sections_dict[i][1](file, table, folder)
                            elif i == 15:
                                sections_dict[i][1](file, table, product_type)
                            else:
                                sections_dict[i][1](file, table)

                else:
                    continue  # If the file is empty, move to the next one
            else:
                continue  # If the JSON file is missing, move to the next section
        dt_now = datetime.datetime.now()  # Get the current date
        dt_now = dt_now.strftime("%Y_%m_%d-%H_%M_%S")
        file_name = f'{folder}\\reports\\{computer}_{dt_now}.docx'
        doc_template.save(file_name)  # Save the report to a file
        return file_name

    except Exception:  # Log any errors
        with open(f'{folder}\\errors\\python_errors.log', 'a') as f:
            f.write(f'{computer}\nОшибка при заполнении Секции {i}:\n')
            f.write('{}\n'.format(traceback.format_exc()))
        return 'Error'


if __name__ == "__main__":
    args = parse_arguments()
    print(report_create(args.computer, args.path, args.report_version, args.os_version, args.ratio))
